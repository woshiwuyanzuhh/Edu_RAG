"""Word 文档解析器 — python-docx。"""

from docx import Document as DocxDocument

from src.interfaces.parser import IParser
from src.shared.exceptions import ParseError


class DocxParser(IParser):
    @property
    def supported_extensions(self) -> set[str]:
        return {".docx", ".doc"}

    def parse(self, file_path: str) -> str:
        try:
            doc = DocxDocument(file_path)
        except Exception as e:
            raise ParseError(f"DOCX 解析失败: {file_path}", detail=str(e))

        parts = []

        # 段落
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        # 表格内容（解决表格被丢弃的问题）
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts).strip()
